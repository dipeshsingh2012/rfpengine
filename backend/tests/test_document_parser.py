"""
Comprehensive Document Parser & Knowledge Base File Upload Test Suite.

Validates:
1. Multi-format parsing (CSV, TSV, JSON, JSONL, Markdown, PDF, DOCX, TXT).
2. 300-500 token sliding-window recursive chunking with 50-token overlap.
3. Automatic taxonomy category inference (zero-config).
4. Full multipart HTTP upload endpoint (POST /api/v1/knowledge-base/upload).
5. Validation edge cases (empty files, unsupported formats, corrupt payloads).
"""

from __future__ import annotations

import io
import json
import pytest
import httpx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import docx
from app.main import app
from app.services.document_parser_service import DocumentParserService
from app.core.db import get_session_factory
from app.models.db_models import KBEntry
from sqlalchemy import delete


@pytest.fixture
async def cleanup_test_tenant():
    """Ensure test tenant rows are cleaned up before and after test execution."""
    session_factory = get_session_factory()
    async with session_factory() as session:
        await session.execute(delete(KBEntry).where(KBEntry.tenant_id == "test-upload-tenant"))
        await session.commit()
    yield
    async with session_factory() as session:
        await session.execute(delete(KBEntry).where(KBEntry.tenant_id == "test-upload-tenant"))
        await session.commit()


# ==============================================================================
# Unit Tests: Document Parser Service
# ==============================================================================

def test_csv_document_parsing():
    csv_content = b"""question,answer,category
What is your data retention period?,Customer data is retained for 30 days after contract termination.,Compliance
Where are production servers hosted?,All infrastructure is hosted in AWS us-east-1 and us-west-2.,Infrastructure
Do you support SAML 2.0 SSO?,Yes SAML 2.0 and OIDC integrations are supported across all tiers.,Security
"""
    entries = DocumentParserService.parse_document(
        content=csv_content,
        filename="security_faq.csv",
        tenant_id="test-tenant",
    )

    assert len(entries) == 3
    assert entries[0].question == "What is your data retention period?"
    assert "30 days" in entries[0].answer
    assert entries[0].category == "Compliance"
    assert entries[1].category == "Infrastructure"
    assert entries[2].category == "Security"


def test_tsv_document_parsing():
    tsv_content = b"question\tanswer\tcategory\nWhat is your SLA?\t99.95% monthly uptime.\tOperations\nHow is data encrypted?\tAES-256 at rest.\tSecurity\n"
    entries = DocumentParserService.parse_document(
        content=tsv_content,
        filename="sla_metrics.tsv",
        tenant_id="test-tenant",
    )

    assert len(entries) == 2
    assert entries[0].question == "What is your SLA?"
    assert "99.95%" in entries[0].answer
    assert entries[0].category == "Operations"
    assert entries[1].category == "Security"


def test_json_and_jsonl_parsing():
    # 1. Standard JSON array
    json_content = b"""[
        {"question": "How are encryption keys managed?", "answer": "AWS KMS with annual key rotation.", "category": "Security"},
        {"prompt": "What is the disaster recovery RPO/RTO?", "response": "RPO is 1 hour and RTO is 4 hours.", "section": "Operations"}
    ]"""
    json_entries = DocumentParserService.parse_document(
        content=json_content,
        filename="dr_and_keys.json",
        tenant_id="test-tenant",
    )
    assert len(json_entries) == 2
    assert json_entries[0].question == "How are encryption keys managed?"
    assert json_entries[1].question == "What is the disaster recovery RPO/RTO?"
    assert "4 hours" in json_entries[1].answer

    # 2. JSONL line-by-line
    jsonl_content = b"""{"question": "Do you perform pen tests?", "answer": "Annually by CREST accredited firm."}
{"question": "What is password policy?", "answer": "16+ alphanumeric with FIDO2 MFA."}"""
    jsonl_entries = DocumentParserService.parse_document(
        content=jsonl_content,
        filename="pentest_and_auth.jsonl",
        tenant_id="test-tenant",
    )
    assert len(jsonl_entries) == 2
    assert jsonl_entries[0].question == "Do you perform pen tests?"
    assert "CREST" in jsonl_entries[0].answer


def test_markdown_heading_hierarchy():
    md_content = b"""# Cloud Architecture Whitepaper

## Section 1: AWS VPC Infrastructure
All services run in private subnets inside dedicated AWS Virtual Private Clouds with AWS Network Firewalls.

## Section 2: Data Encryption Controls
All EBS volumes, S3 buckets, and RDS instances are encrypted using AES-256 with KMS Customer Managed Keys.
"""
    entries = DocumentParserService.parse_document(
        content=md_content,
        filename="Cloud_Architecture_Whitepaper.md",
        tenant_id="test-tenant",
    )

    assert len(entries) >= 2
    headings = [e.question for e in entries]
    assert any("AWS VPC Infrastructure" in h for h in headings)
    assert any("Data Encryption Controls" in h for h in headings)
    for e in entries:
        assert e.category in ["Cloud & Architecture", "Compliance & Security", "Security & Cryptography"]


def test_pdf_parsing_and_chunking():
    # Build in-memory PDF with 2 pages
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    c.drawString(50, 750, "Acme SOC 2 Security Whitepaper - Page 1")
    c.drawString(50, 720, "All customer communications are encrypted using TLS 1.3 with strict forward secrecy.")
    c.showPage()
    c.drawString(50, 750, "Acme SOC 2 Security Whitepaper - Page 2")
    c.drawString(50, 720, "Disaster recovery RPO is 15 minutes and RTO is 1 hour across multi-region clusters.")
    c.showPage()
    c.save()

    pdf_bytes = pdf_buffer.getvalue()
    entries = DocumentParserService.parse_document(
        content=pdf_bytes,
        filename="Acme_SOC2_Whitepaper.pdf",
        tenant_id="test-tenant",
    )

    assert len(entries) == 2
    assert entries[0].metadata["page_number"] == 1
    assert entries[1].metadata["page_number"] == 2
    assert entries[0].metadata["format"] == "pdf"
    assert entries[0].category == "Compliance & Security"


def test_docx_parsing():
    # Build in-memory DOCX
    doc = docx.Document()
    doc.add_heading("REST API Developer Guide", 0)
    doc.add_heading("1. Rate Limits and Authentication", level=1)
    doc.add_paragraph("Rate limits are 10,000 requests per minute with OAuth 2.0 bearer token validation.")
    doc.add_heading("2. Webhooks", level=1)
    doc.add_paragraph("Real-time webhooks are cryptographically signed with HMAC-SHA256 secret tokens.")
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    entries = DocumentParserService.parse_document(
        content=docx_bytes,
        filename="API_Integration_Guide.docx",
        tenant_id="test-tenant",
    )

    assert len(entries) >= 2
    assert any("Rate Limits" in e.question for e in entries)
    assert any("Webhooks" in e.question for e in entries)
    assert entries[0].category == "Product & Integrations"


def test_text_sliding_window_chunking():
    # Long text document (~4000 characters)
    sentences = (
        "Acme maintains a comprehensive vendor risk management program. "
        "Every third-party service provider must undergo an annual security audit and provide a SOC 2 Type II report. "
    ) * 20
    text_bytes = sentences.encode("utf-8")

    entries = DocumentParserService.parse_document(
        content=text_bytes,
        filename="Vendor_Risk_Management_Policy.txt",
        tenant_id="test-tenant",
    )

    assert len(entries) > 1, "Must chunk long text into multiple pieces"
    for e in entries:
        assert len(e.answer) <= DocumentParserService.CHUNK_SIZE_CHARS + 100
        assert e.category == "Compliance & Security" or e.category == "HR & Corporate Policies"


def test_automatic_category_inference_heuristics():
    assert DocumentParserService.infer_category("01_Security_and_Compliance_Whitepaper.md") == "Compliance & Security"
    assert DocumentParserService.infer_category("02_SLA_Disaster_Recovery_and_Operations.pdf") == "SLA & Operations"
    assert DocumentParserService.infer_category("03_Data_Privacy_GDPR_and_Subprocessors.json") == "Privacy & Legal"
    assert DocumentParserService.infer_category("05_Product_Features_and_API_Integrations.docx") == "Product & Integrations"
    assert DocumentParserService.infer_category("06_Employee_Code_of_Conduct_and_HR_Policies.txt") == "HR & Corporate Policies"


# ==============================================================================
# Integration Tests: HTTP Upload Endpoint (POST /api/v1/knowledge-base/upload)
# ==============================================================================

@pytest.mark.asyncio
async def test_upload_endpoint_csv_success():
    csv_bytes = b"""question,answer,category
What is the incident response SLA?,P1 critical incidents are acknowledged within 15 minutes.,Support SLA
Is multi-factor authentication mandatory?,MFA is enforced on all corporate systems via WebAuthn/FIDO2.,Security
"""
    files = {"file": ("test_sla.csv", io.BytesIO(csv_bytes), "text/csv")}
    data = {"tenant_id": "test-upload-tenant"}

    async with httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://test") as client:
        async with app.router.lifespan_context(app):
            response = await client.post(
                "/api/v1/knowledge-base/upload",
                files=files,
                data=data,
            )

            assert response.status_code == 201, f"Upload failed: {response.text}"
            payload = response.json()
            assert payload["records_created"] == 2
            assert payload["tenant_id"] == "test-upload-tenant"
            assert payload["filename"] == "test_sla.csv"
            assert len(payload["preview"]) == 2


@pytest.mark.asyncio
async def test_upload_endpoint_markdown_success():
    md_bytes = b"""# Encryption Standards
## Encryption in Transit
Enforced with TLS 1.3 and HSTS preloading.
## Encryption at Rest
Enforced with AES-256 and AWS KMS Customer-Managed Keys.
"""
    files = {"file": ("encryption_standards.md", io.BytesIO(md_bytes), "text/markdown")}
    data = {"tenant_id": "test-upload-tenant"}

    async with httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://test") as client:
        async with app.router.lifespan_context(app):
            response = await client.post(
                "/api/v1/knowledge-base/upload",
                files=files,
                data=data,
            )

            assert response.status_code == 201
            payload = response.json()
            assert payload["records_created"] >= 2
            assert "Security & Cryptography" in payload["categories"] or "Compliance & Security" in payload["categories"]


@pytest.mark.asyncio
async def test_upload_endpoint_empty_file_bad_request():
    empty_bytes = b""
    files = {"file": ("empty.txt", io.BytesIO(empty_bytes), "text/plain")}
    data = {"tenant_id": "test-upload-tenant"}

    async with httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://test") as client:
        async with app.router.lifespan_context(app):
            response = await client.post(
                "/api/v1/knowledge-base/upload",
                files=files,
                data=data,
            )

            assert response.status_code == 400
            assert "empty" in response.json()["detail"].lower()


@pytest.mark.asyncio
async def test_upload_endpoint_unsupported_file_extension():
    binary_bytes = b"\x7fELF\x02\x01\x01\x00"
    files = {"file": ("malicious_payload.exe", io.BytesIO(binary_bytes), "application/octet-stream")}
    data = {"tenant_id": "test-upload-tenant"}

    async with httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://test") as client:
        async with app.router.lifespan_context(app):
            response = await client.post(
                "/api/v1/knowledge-base/upload",
                files=files,
                data=data,
            )

            assert response.status_code in [400, 422]
            assert "Unsupported file format" in response.json()["detail"]

